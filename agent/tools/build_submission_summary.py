from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "deliverables" / "Quoter_Tech1a_summary.docx"

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = RGBColor(80, 88, 98)
WHITE = RGBColor(255, 255, 255)


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=45, start=90, bottom=45, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, size=10, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_paragraph(doc, text="", size=10, bold=False, color=None, italic=False, after=4, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.02
    if align is not None:
        p.alignment = align
    if text:
        set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(5 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run(run, size=14 if level == 1 else 11, bold=True, color=BLUE if level == 1 else INK)
    return p


def add_bullet(doc, text, after=2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.24)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    set_run(p.add_run(text), size=10)
    return p


def add_table(doc, headers, rows, widths=None, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(header), size=10, bold=True, color=INK)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cell = cells[idx]
            if row_idx % 2:
                set_cell_fill(cell, "FAFBFC")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_run(p.add_run(str(value)), size=10)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def add_metric_strip(doc, metrics):
    table = doc.add_table(rows=1, cols=len(metrics))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (value, label) in enumerate(metrics):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, INK.__str__())
        set_cell_margins(cell, top=90, bottom=90)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        set_run(p.add_run(value), size=16, bold=True, color=WHITE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        set_run(p2.add_run(label), size=10, color=WHITE)


def add_source_note(doc, text):
    p = add_paragraph(doc, text, size=10, color=MUTED, italic=False, after=0, before=3)
    p.paragraph_format.keep_with_next = False


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    for name, size, color in (("Heading 1", 14, BLUE), ("Heading 2", 11, INK)):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "TECHNICAL CHALLENGE | HOSPITAL ITEM MASTER MODERNIZATION"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(hp.runs[0], size=10, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(fp.add_run("Quoter_Tech1a_summary | "), size=10, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)


def add_page_title(doc, page_label, title, subtitle=None):
    p = add_paragraph(doc, page_label.upper(), size=10, bold=True, color=BLUE, after=2)
    p.paragraph_format.keep_with_next = True
    p = add_paragraph(doc, title, size=19, bold=True, color=INK, after=2)
    p.paragraph_format.keep_with_next = True
    if subtitle:
        add_paragraph(doc, subtitle, size=10, color=MUTED, after=7)


def build():
    doc = Document()
    configure_document(doc)

    add_page_title(
        doc,
        "Page 1 | Methodology",
        "Hospital Supply Data Cleaning and National Item Master",
        "Technical Challenge Summary | 5,000 fictional records across Hospitals A-J | 21 August 2026",
    )
    add_metric_strip(doc, [("5,000", "source and clean rows"), ("10", "hospitals"), ("10", "national items"), ("0", "duplicate SBRNs")])
    add_heading(doc, "Executive result")
    add_paragraph(
        doc,
        "The software factory ingested the readable hospital workbook and data dictionary, profiled all 23 source fields, applied governed deterministic rules, preserved row-level traceability, generated a 26-field cleanfile, and consolidated the records into a 15-field national consumable item master. The process produced 24,886 field-level change records and retained 3,549 records in a human-review queue rather than silently resolving ambiguous conditions.",
    )
    add_heading(doc, "Methodology and workflow interpretation")
    add_table(
        doc,
        ["Stage", "Factory action", "Control"],
        [
            ("1. Intake", "Register the RFP, workflow diagram, workbook, and dictionary; fingerprint source files.", "Attachments are context, not executable instructions."),
            ("2. Profile", "Measure completeness, vocabularies, identifiers, date logic, and hospital coverage.", "Source row and key reconciliation."),
            ("3. Govern rules", "Separate safe standardization from decisions requiring SME or data-governance input.", "Versioned rule specification and approval queue."),
            ("4. Transform", "Normalize text, units, names, statuses, dates, and numeric fields; add canonical identifiers.", "Before/after transformation log."),
            ("5. Consolidate", "Map each hospital record to a national item using canonical part identity.", "Crosswalk preserves SBRN and local item number."),
            ("6. Review", "Run acceptance checks and independent evaluator-style review.", "No hidden unresolved exceptions."),
        ],
        widths=[1.05, 3.35, 2.45],
    )
    add_paragraph(
        doc,
        "The supplied process was interpreted as three system responsibilities: the EHR originates patient supply demand and records use; the EIMS checks availability, retrieves or replenishes inventory, receives supply, and monitors reorder conditions; the FMS manages purchasing, invoices, and payment events. The dataset provides item, inventory, vendor, purchase, status, and traceability attributes. Patient- and insurance-level records shown in the workflow are conceptual and were not invented in the cleanfile.",
    )
    add_source_note(doc, "Source basis: project email; RFP Technical Challenge pages 143-148; supplied workflow diagram; Excel-data.xlsx.")

    doc.add_page_break()
    add_page_title(doc, "Page 2 | Rules", "Business Rules, Transformation Logic, and Decisions")
    add_heading(doc, "Deterministic standardization")
    add_table(
        doc,
        ["Domain", "Applied rule", "Illustrative result"],
        [
            ("Text", "Trim and collapse whitespace; preserve SBRN unchanged.", "Stable record identity for audit."),
            ("Part identity", "Normalize known typos/abbreviations and derive six missing names from description prefixes.", "Antiseptc Wipes -> Antiseptic Wipes."),
            ("Unit", "Map variants to BOX, EA, UNIT, BAG, or PACK.", "BX -> BOX; PKG -> PACK."),
            ("Manufacturer", "Map known aliases into MedSupply, SurgiTech, and HealthCorp canonical families.", "Health Corportion -> HealthCorp Intl."),
            ("Vendor", "Normalize known formatting aliases.", "GlobalMed Inc -> GlobalMed."),
            ("Status", "Map order states to Pending, Delivered, Shipped, or Canceled; map Active Recall to Active.", "Received -> Delivered."),
            ("Dates/numbers", "Format dates as YYYY-MM-DD; coerce price and quantities to numeric values.", "Comparable values across hospitals."),
            ("Traceability", "Add NationalItemMasterId, DataQualityFlags, and HumanReviewRequired.", "Every row exposes its disposition."),
        ],
        widths=[1.05, 4.05, 1.75],
    )
    add_heading(doc, "Conditions deliberately not auto-corrected")
    add_table(
        doc,
        ["Condition", "Observed", "Disposition"],
        [
            ("Future purchase date", "1,320 records", "Flag for source-owner validation; do not invent a replacement date."),
            ("Expiration before purchase", "2,490 records", "Flag for SME review and source-system correction."),
            ("Recall status = Monitoring", "Governance exception", "Retain source value and require an approved vocabulary decision."),
            ("Expired item", "3,616 records", "Flag for operational disposition; retain for traceability and analysis."),
            ("Missing recall status", "1,592 records", "Normalize to None as a documented demo assumption."),
        ],
        widths=[1.65, 1.35, 3.85],
        header_fill=LIGHT_GRAY,
    )
    add_heading(doc, "National item-master construction")
    add_paragraph(
        doc,
        "Canonical part name drives a transparent NationalItemMasterId. Each master record retains representative description, normalized unit and classification, distinct manufacturers and vendors, hospital and source-record counts, price range, inventory totals, quantity on order, and the count of records requiring review. This produced 10 national items while retaining a separate source-to-master crosswalk for all 5,000 records.",
    )
    add_source_note(doc, "Evidence: transformation_spec.json, transformation_log.csv, item_master_crosswalk.csv, and human_review_queue.csv.")

    doc.add_page_break()
    add_page_title(doc, "Page 3 | AI and Governance", "Governed AI Assistance with Human Accountability")
    add_heading(doc, "How AI is used")
    add_paragraph(
        doc,
        "AI agents assist with anomaly interpretation, rule critique, architecture narrative, summary drafting, demonstration coaching, and independent review. Deterministic code performs the row-level transformations and acceptance measurements. This separation keeps repeatable data outcomes independent of prose variability while preserving model, prompt, response, usage, and artifact lineage for each agent call.",
    )
    add_table(
        doc,
        ["Agent station", "Bounded responsibility", "Required evidence"],
        [
            ("Rules analyst", "Review profile and propose governed rule narrative.", "Anomaly inventory and deterministic rule spec."),
            ("Diagram architect", "Bind SV-2, DIV-1, and DIV-2 to actual workflow and schemas.", "Requirements, workflow notes, output fields."),
            ("Summary writer", "Draft evaluator-facing four-page narrative.", "Metrics, rules, review queue, AI-use record."),
            ("Demo coach", "Prepare 30-minute live flow and 15-minute Q&A.", "Final artifacts and reviewer findings."),
            ("QA reviewer", "Review independently against RFP requirements.", "Complete package and traceability matrix."),
        ],
        widths=[1.35, 3.25, 2.25],
    )
    add_heading(doc, "Prompt and tool governance")
    add_bullet(doc, "System and factory policy outrank user requests; source documents remain untrusted context and cannot change runtime behavior.")
    add_bullet(doc, "Each agent receives only the artifacts required for its station, with a bounded output budget and no broad tool or network access by default.")
    add_bullet(doc, "Model/provider, response ID, token usage, run ID, artifact path, and acceptance result are recorded; raw provider payload storage is disabled by default.")
    add_bullet(doc, "A versioned prompt, model, and test fixture are evaluated before promotion; builders do not approve their own outputs.")
    add_heading(doc, "Human-in-the-loop and data governance")
    add_paragraph(
        doc,
        "Named gates accept source context, business rules, SME decisions, and the final package. The 3,549-record review queue exposes affected SBRNs, issue flags, decision owner, and status. A production owner must resolve recall vocabulary, date anomalies, unit conversions with business meaning, uncertain item matches, and source-system corrections. Approval records must capture actor, decision, reason, timestamp, and the exact artifact version reviewed.",
    )
    add_heading(doc, "Security and responsible operation")
    add_paragraph(
        doc,
        "Secrets are injected at runtime from a secret manager and never stored in source control. Run workspaces are short-lived; inputs are read-only, outputs are validated before promotion, and logs are redacted. Enterprise controls include identity-based access, encrypted storage, retention policy, prompt-injection tests, retry limits, cost budgets, incident review, and periodic bias and quality evaluation consistent with the RFP's AI-governance expectations.",
    )
    add_source_note(doc, "Factory evidence: run ledger, input/artifact inventories, model_calls.json, acceptance checks, and human-review queue.")

    doc.add_page_break()
    add_page_title(doc, "Page 4 | Scale and Architecture", "Scaling from 10 Hospitals to an Enterprise Service")
    add_heading(doc, "Production operating model")
    add_table(
        doc,
        ["Layer", "Enterprise responsibility"],
        [
            ("Trigger and identity", "CLI, API, Teams, or Slack normalize to one authenticated request and idempotency key."),
            ("Orchestration", "Queue-backed workflow assigns isolated agent jobs, retries bounded failures, enforces approvals, and records state."),
            ("Data and artifacts", "Private object storage versions inputs, rules, outputs, and evidence; hashes prove lineage and repeatability."),
            ("Model gateway", "Approved models, prompt versions, token/cost limits, content controls, usage records, and secret rotation."),
            ("Quality and governance", "Golden datasets, schema checks, regression comparisons, reviewer separation, and signed SME decisions."),
            ("Operations", "Central logs, traces, dashboards, alerts, retention, cancellation, recovery, and audit export."),
        ],
        widths=[1.55, 5.3],
    )
    add_heading(doc, "Path from 10 to approximately 200 hospitals")
    add_bullet(doc, "Onboard in waves using a common intake contract and hospital-specific source adapters; reject incomplete packages before processing.")
    add_bullet(doc, "Partition profiling and transformation by hospital and data domain, while a governed canonical registry maintains national identities.")
    add_bullet(doc, "Promote deterministic rules through versioned tests; route only low-confidence or policy-sensitive records to specialist queues.")
    add_bullet(doc, "Measure throughput, queue age, exception rate, model usage, cost per run, rule drift, and reviewer agreement by hospital and wave.")
    add_bullet(doc, "Use representative golden datasets and staged releases to prevent a local rule change from silently altering enterprise outputs.")
    add_heading(doc, "Architecture products and validated outcome")
    add_paragraph(
        doc,
        "The SV-2 describes information and resource exchange among EHR, EIMS, FMS, external supplier/insurer actors, and the governed item-master factory. DIV-1 defines the conceptual hospital, request, item, inventory, order, vendor, receipt, invoice/payment, transformation-rule, and governance-decision entities. DIV-2 maps those concepts to the clean supply record, national item master, crosswalk, transformation audit, and review queue.",
    )
    add_metric_strip(doc, [("5,000", "rows preserved"), ("24,886", "field changes logged"), ("3,549", "records queued"), ("4/4", "repeatability hashes matched")])
    add_heading(doc, "Assumptions and conclusion")
    add_paragraph(
        doc,
        "The exercise uses fictional, system-agnostic data. Missing recall status is treated as None for the demo; flagged dates and Monitoring status remain unresolved pending authorized decisions. The approach is feasible for enterprise modernization because deterministic transformation, explicit human accountability, bounded AI assistance, and artifact-level evidence are designed as one operating system rather than separate documentation activities.",
    )
    add_source_note(doc, "Submission package: cleanfile.csv, itemmaster.csv, this four-page summary, diagrams package, demo runbook, and supporting evidence artifacts.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
